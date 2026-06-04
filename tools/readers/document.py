"""tools/readers/document.py — Extract text and metadata from documents."""
from __future__ import annotations

from pathlib import Path
from typing import Any


def read_document(file_path: str) -> dict[str, Any]:
    """
    Read PDF, DOCX, plain text, Markdown, HTML.
    Returns text preview (first 2,000 chars) and metadata.
    """
    path = Path(file_path)
    ext = path.suffix.lower()
    size_bytes = path.stat().st_size

    metadata: dict[str, Any] = {"size_bytes": size_bytes, "extension": ext}
    text_preview = ""
    page_count = None

    try:
        if ext == ".pdf":
            import fitz  # PyMuPDF

            doc = fitz.open(str(path))
            page_count = doc.page_count
            text_preview = "".join(doc[i].get_text() for i in range(min(3, page_count)))[:2000]
            metadata["page_count"] = page_count
            doc.close()

        elif ext == ".docx":
            import docx

            document = docx.Document(str(path))
            full_text = "\n".join(p.text for p in document.paragraphs)
            text_preview = full_text[:2000]
            metadata["paragraph_count"] = len(document.paragraphs)

        elif ext in (".txt", ".md"):
            with open(path, encoding="utf-8", errors="replace") as f:
                content = f.read()
            text_preview = content[:2000]
            metadata["char_count"] = len(content)

        elif ext in (".html", ".htm"):
            from html.parser import HTMLParser

            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.parts: list[str] = []

                def handle_data(self, data: str) -> None:
                    self.parts.append(data)

            with open(path, encoding="utf-8", errors="replace") as f:
                raw = f.read()
            parser = TextExtractor()
            parser.feed(raw)
            full_text = " ".join(parser.parts)
            text_preview = full_text[:2000]

        else:
            with open(path, encoding="utf-8", errors="replace") as f:
                text_preview = f.read(2000)

    except Exception as exc:
        metadata["read_error"] = str(exc)

    schema = {
        "type": "document_text",
        "text_preview_chars": len(text_preview),
        "page_count": page_count,
    }

    return {
        "schema": schema,
        "sample": text_preview,
        "metadata": metadata,
    }
